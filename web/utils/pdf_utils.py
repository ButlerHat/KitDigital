# Let's redeclare all necessary variables and recreate the image with appropriate left and right margins.
from fpdf import FPDF
from PIL import Image
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_BREAK


def convert_img_to_pdf(image_path: str, output_path: str):
    # A4 dimensions in pixels at 300 DPI for higher resolution
    a4_width_px = 2480  # Approximately 210mm in inches
    a4_height_px = 3508  # Approximately 297mm in inches

    # Margins in pixels (5% of the A4 width on each side)
    margin_px = a4_width_px * 0.05

    # Open the original image to calculate its aspect ratio
    with Image.open(image_path) as img:
        # Calculate the new image dimensions to maintain aspect ratio within the margins
        img_ratio = img.width / img.height
        new_width = a4_width_px - 2 * margin_px
        new_height = new_width / img_ratio

        # Check if the new height is not greater than A4 height
        if new_height > (a4_height_px - 2 * margin_px):
            new_height = a4_height_px - 2 * margin_px
            new_width = new_height * img_ratio

        # Resize the image with high quality resampling filter
        img_resized = img.resize((int(new_width), int(new_height)), Image.LANCZOS)

        # Create a blank A4 white image
        a4_image = Image.new('RGB', (a4_width_px, a4_height_px), 'white')

        # Calculate the position to paste the resized image on the A4 white image to center it
        x_position = int(margin_px)
        y_position = (a4_height_px - int(new_height)) // 4

        # Paste the resized image onto the white background image
        a4_image.paste(img_resized, (x_position, y_position))

        # Save the new image with margins
        output_tmp_path = "".join(image_path.split(".")[:-1]) + "_tmp." + image_path.split(".")[-1]
        a4_image.save(output_tmp_path, format='PNG')

    # Create a PDF document and add the new image with margins
    pdf_with_margins = FPDF(unit="pt", format=(a4_width_px, a4_height_px))
    pdf_with_margins.add_page()
    pdf_with_margins.image(output_tmp_path, 0, 0)

    # Save the PDF file with margins
    pdf_with_margins.output(output_path)


def append_text_and_picture_to_document(filename, identificador, text, picture, width_inches: float=6.0):
    doc = Document(filename)
    
    if not identificador:
        doc.add_paragraph()
        if text:
            doc.add_paragraph(text)
        if picture:
            doc.add_picture(picture, width=Inches(width_inches))
        doc.save(filename)
        return

    finded = False
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if identificador in run.text:
                # Split the text at the identifier
                pre_text, post_text = run.text.split(identificador, 1)
                
                # Update the current run with the text before the identifier
                run.text = pre_text
                
                # If there is a picture to add, insert it in a new run
                if picture:
                    # Adding a new paragraph for the picture
                    pic_paragraph = paragraph.insert_paragraph_before()
                    if text:
                        pic_run = pic_paragraph.add_run(text)
                        pic_run.add_break(WD_BREAK.LINE)
                    else:
                        pic_run = pic_paragraph.add_run()
                    pic_run.add_picture(picture, width=Inches(width_inches))
                # Code done with chatgpt and bullshit, but works
                else:
                    if text:
                        pic_paragraph = paragraph.insert_paragraph_before()
                        if text:
                            pic_run = pic_paragraph.add_run(text)
                            pic_run.add_break(WD_BREAK.LINE)
                
                # Add a new run after the picture with the identifier and any text following it
                paragraph.add_run(identificador + post_text)
                
                finded = True
                break  # Asumimos que hay un solo identificador por párrafo
        
        if finded:
            break  # No need to check other paragraphs once found

    if not finded:
        raise AssertionError(f"No se ha encontrado el identificador {identificador} en el documento {filename}")

    doc.save(filename)


def remove_identifier(filename, identificador):
    """
    Remove text from word document.
    """
    doc = Document(filename)
    
    finded = False
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if identificador in run.text:
                # Split the text at the identifier
                pre_text, post_text = run.text.split(identificador, 1)
                
                # Update the current run with the text before the identifier
                run.text = pre_text + post_text
                
                finded = True
                break  # Asumimos que hay un solo identificador por párrafo
        
        if finded:
            break  # No need to check other paragraphs once found

    if not finded:
        raise AssertionError(f"No se ha encontrado el identificador {identificador} en el documento {filename}")

    doc.save(filename)


def check_if_identifier_exists(filename, identificador):
    """
    Check if identifier exists in word document.
    """
    doc = Document(filename)
    
    finded = False
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if identificador in run.text:
                finded = True
                break  # Asumimos que hay un solo identificador por párrafo
        
        if finded:
            break  # No need to check other paragraphs once found

    return finded


if __name__ == "__main__":
    # Redefine the original image path
    original_image_path = '/workspaces/ai-butlerhat/data-butlerhat/robotframework-butlerhat/TestSuites/KitDigital/result_kit/djadelpeluqueriaes/logo_kit_digital/logo_screenshot.png'
    new_image_path_with_margins = '/workspaces/ai-butlerhat/data-butlerhat/robotframework-butlerhat/TestSuites/KitDigital/result_kit/djadelpeluqueriaes/logo_kit_digital/logo_screenshot_edited.png'
    # Convert the image to PDF
    convert_img_to_pdf(original_image_path, new_image_path_with_margins)
