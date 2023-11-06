from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_BREAK

def create_new_document(filename):
    doc = Document()
    doc.save(filename)



def append_text_and_picture_to_document(filename, identificador, text, picture, width_inches: float=6.0):
    doc = Document(filename)
    
    if not identificador:
        doc.add_paragraph()
        if text:
            doc.add_paragraph(text)
        if picture:
            doc.add_picture(picture, width=Inches(width_inches))
        doc.save(filename)
        return

    finded = False
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if identificador in run.text:
                # Split the text at the identifier
                pre_text, post_text = run.text.split(identificador, 1)
                
                # Update the current run with the text before the identifier
                run.text = pre_text
                
                # If there is a picture to add, insert it in a new run
                if picture:
                    # Adding a new paragraph for the picture
                    pic_paragraph = paragraph.insert_paragraph_before()
                    if text:
                        pic_run = pic_paragraph.add_run(text)
                        pic_run.add_break(WD_BREAK.LINE)
                    else:
                        pic_run = pic_paragraph.add_run()
                    pic_run.add_picture(picture, width=Inches(width_inches))
                else:
                    if text:
                        run.add_run(text)
                        run.add_break(WD_BREAK.LINE)
                
                # Add a new run after the picture with the identifier and any text following it
                paragraph.add_run(identificador + post_text)
                
                finded = True
                break  # Asumimos que hay un solo identificador por párrafo
        
        if finded:
            break  # No need to check other paragraphs once found

    if not finded:
        raise AssertionError(f"No se ha encontrado el identificador {identificador} en el documento {filename}")

    doc.save(filename)


def remove_identifier(filename, identificador):
    """
    Remove text from word document.
    """
    doc = Document(filename)
    
    finded = False
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if identificador in run.text:
                # Split the text at the identifier
                pre_text, post_text = run.text.split(identificador, 1)
                
                # Update the current run with the text before the identifier
                run.text = pre_text + post_text
                
                finded = True
                break  # Asumimos que hay un solo identificador por párrafo
        
        if finded:
            break  # No need to check other paragraphs once found

    if not finded:
        raise AssertionError(f"No se ha encontrado el identificador {identificador} en el documento {filename}")

    doc.save(filename)


def check_if_identifier_exists(filename, identificador):
    """
    Check if identifier exists in word document.
    """
    doc = Document(filename)
    
    finded = False
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if identificador in run.text:
                finded = True
                break  # Asumimos que hay un solo identificador por párrafo
        
        if finded:
            break  # No need to check other paragraphs once found

    return finded