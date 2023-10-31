from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_BREAK

def create_new_document(filename):
    doc = Document()
    doc.save(filename)



def append_text_and_picture_to_document(filename, text, picture):
    doc = Document(filename)
    # Añadir texto y luego un salto de línea
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.add_break(WD_BREAK.LINE)

    # Añadir la imagen
    run = paragraph.add_run()
    run.add_picture(picture, width=Inches(6.0))
    doc.save(filename)
